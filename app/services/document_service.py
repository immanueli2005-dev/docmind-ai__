# app/services/document_service.py - Text extraction and chunking services

import os
import docx
from pypdf import PdfReader
from typing import List, Dict, Any

def extract_document_pages(filepath: str, file_type: str) -> List[Dict[str, Any]]:
    """Extract page-by-page text content from file path and return list of dicts with page indices."""
    if file_type == "pdf":
        return _extract_from_pdf(filepath)
    elif file_type == "docx":
        return _extract_from_docx(filepath)
    elif file_type == "txt":
        return _extract_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def _extract_from_pdf(filepath: str) -> List[Dict[str, Any]]:
    pages_content = []
    try:
        reader = PdfReader(filepath)
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            # Clean consecutive whitespace
            cleaned_text = " ".join(text.split())
            if cleaned_text.strip():
                pages_content.append({
                    "page": page_num + 1,
                    "content": cleaned_text
                })
    except Exception as e:
        print(f"Error reading PDF {filepath}: {e}")
    return pages_content

def _extract_from_docx(filepath: str) -> List[Dict[str, Any]]:
    pages_content = []
    try:
        doc = docx.Document(filepath)
        current_paragraphs = []
        page_num = 1
        
        # Group paragraphs to simulate "pages" for DOCX documents
        for i, para in enumerate(doc.paragraphs):
            cleaned_para = " ".join(para.text.split())
            if cleaned_para.strip():
                current_paragraphs.append(cleaned_para)
            
            # Flush grouping every 4 paragraphs or at the end of paragraphs list
            if (len(current_paragraphs) >= 4) or (i == len(doc.paragraphs) - 1):
                if current_paragraphs:
                    pages_content.append({
                        "page": page_num,
                        "content": "\n".join(current_paragraphs)
                    })
                    current_paragraphs = []
                    page_num += 1
    except Exception as e:
        print(f"Error reading DOCX {filepath}: {e}")
    return pages_content

def _extract_from_txt(filepath: str) -> List[Dict[str, Any]]:
    pages_content = []
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
            
        # Split text into virtual pages of ~1200 characters each
        char_limit = 1200
        page_num = 1
        for start in range(0, len(text), char_limit):
            end = min(start + char_limit, len(text))
            cleaned_segment = " ".join(text[start:end].split())
            if cleaned_segment.strip():
                pages_content.append({
                    "page": page_num,
                    "content": cleaned_segment
                })
                page_num += 1
    except Exception as e:
        print(f"Error reading TXT {filepath}: {e}")
    return pages_content

def chunk_document_pages(pages_content: List[Dict[str, Any]], chunk_size: int = 500, chunk_overlap: int = 100) -> List[Dict[str, Any]]:
    """Segment extracted pages into overlapping text chunks with page indices and paragraph counts."""
    chunks = []
    
    for page_obj in pages_content:
        page_num = page_obj["page"]
        text = page_obj["content"]
        
        words = text.split()
        if not words:
            continue
            
        # Approximate word limits (average 6 chars per word)
        word_limit = chunk_size // 6
        overlap_words = chunk_overlap // 6
        
        start_idx = 0
        para_counter = 1
        
        while start_idx < len(words):
            end_idx = min(start_idx + word_limit, len(words))
            chunk_content = " ".join(words[start_idx:end_idx])
            
            chunks.append({
                "page": page_num,
                "paragraph": para_counter,
                "content": chunk_content
            })
            
            para_counter += 1
            start_idx += max(1, word_limit - overlap_words)
            
    return chunks
